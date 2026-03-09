# polara_checker/extraction.py

from __future__ import annotations

import re
import unicodedata
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


# Custom error classes because callers can catch ExtractionError specifically the API endpoint can return a 422 to the user instead of crashing with a 500
class ExtractionError(Exception):
    """
    Raised when a file can't be extracted at all — corrupt file,
    unsupported format, that kind of thing. Not for minor cleanup issues.
    """
    pass

class LowQualityExtractionWarning(Warning):
    """
    Raised as a warning (not an error) when extraction works but the
    output is suspiciously short or garbled. The pipeline still runs,
    but something upstream might be wrong.
    """
    pass


def extract_text(file_path: str | Path) -> str:
    """
    Extract clean plaintext from an uploaded file.

    Supports: .pdf, .docx, .doc, .txt, .md

    Returns clean UTF-8 text. Raises ExtractionError if it can't.
    """

    path = Path(file_path)

    if not path.exists():
        raise ExtractionError(f"File not found: {file_path}")
    
    suffix = path.suffix.lower()

    # Map file extensions to extractor functions
    extractors = {
        ".pdf":  _extract_pdf,
        ".docx": _extract_docx,
        ".doc":  _extract_docx,   # python-docx handles many .doc files too
        ".txt":  _extract_text_file,
        ".md":   _extract_text_file,
    }

    if suffix not in extractors:
        raise ExtractionError(
            f"Unsupported file format: {suffix}"
            f"Supported: {', '.join(extractors)}"
            )
    
    raw = extractors[suffix](path)
    cleaned = _clean(raw)

    # Warn if the result is suspiciosly short
    if len(cleaned.strip()) < 50:
        import warnings
        warnings.warn(
            f"Extracted text from '{path.name}' is only {len(cleaned)} characters. "
            "Extraction may have failed silently.",
            LowQualityExtractionWarning,
            stacklevel=2,
        )
    
    return cleaned

def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        raise ExtractionError(f"Could not open PDF '{path.name}': {exc}") from exc
    
    pages = []
    for page in doc:
        text = page.get_text("text")
        pages.append(text)
    
    doc.close()
    raw = "\n\n".join(pages)

    return raw

def _extract_docx(path: Path) -> str:
    """Extract text from a Word document using python-docx."""
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ExtractionError(f"Could not open Word document '{path.name}': {exc}") from exc
    
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # For tables go row by row, joined with pipe seperators so structure is preserved
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)

def _extract_text_file(path: Path) -> str:
    """Read a plain text or Markdown file."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError(f"Could not decode '{path.name}' with any known encoding.")

# These patterns get stripped from extracted text.
# Each tuple: (compiled regex, replacement)
_STRIP_PATTERNS = [
    # Standalone page numbers: "1", "Page 3", "2 of 10"
    (re.compile(r"^\s*(page\s+)?\d+(\s+of\s+\d+)?\s*$", re.IGNORECASE | re.MULTILINE), ""),

    # Confidentiality footers that repeat on every page
    (re.compile(
        r"(confidential|internal use only|proprietary|do not distribute)[^\n]*",
        re.IGNORECASE
    ), ""),

    # Short all-caps lines (headers/footers like "ACME CORP — INTERNAL")
    (re.compile(r"^[A-Z\s\-–—|]{4,40}$", re.MULTILINE), ""),

    # Three or more blank lines → two blank lines
    (re.compile(r"\n{3,}"), "\n\n"),

    # Trailing whitespace on each line
    (re.compile(r"[ \t]+$", re.MULTILINE), ""),
    (re.compile(r"^[ \t]+", re.MULTILINE), ""),
]

def _clean(text: str) -> str:
    """
    Clean raw extracted text.

    Steps:
    1. Unicode normalization — fixes smart quotes, ligatures like 'fi', etc.
    2. Remove non-printable control characters (but keep newlines and tabs)
    3. Apply the strip patterns above
    4. Final trim
    """
    if not text:
        return ""

    # Step 1: NFC normalization. Combines characters and fixes ligatures.
    text = unicodedata.normalize("NFC", text)

    # Step 2: Remove control characters (binary garbage from bad PDFs)
    text = "".join(
        ch for ch in text
        if unicodedata.category(ch) not in ("Cc", "Cf") or ch in "\n\t"
    )

    # Step 3: Apply strip patterns
    for pattern, replacement in _STRIP_PATTERNS:
        text = pattern.sub(replacement, text)

    # Step 4: Final strip
    return text.strip()