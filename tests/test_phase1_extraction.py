# tests/test_phase1_extraction.py

import warnings
from pathlib import Path

import pytest
from docx import Document

from polara_checker.extraction import (
    ExtractionError,
    LowQualityExtractionWarning,
    _clean,
    extract_text,
)


# ---------------------------------------------------------------------------
# _clean() unit tests
# ---------------------------------------------------------------------------

class TestClean:

    def test_strips_standalone_page_numbers(self):
        text = "Important policy text.\n\n1\n\nMore text.\n\n3 of 10"
        result = _clean(text)
        lines = result.splitlines()
        assert "1" not in lines
        assert "3 of 10" not in lines
        assert "Important policy text." in result

    def test_strips_confidentiality_footer(self):
        text = "Confidential - Acme Corp\nActual policy content here."
        result = _clean(text)
        assert "Confidential" not in result
        assert "Actual policy content" in result

    def test_collapses_excessive_blank_lines(self):
        text = "Paragraph one.\n\n\n\n\nParagraph two."
        result = _clean(text)
        assert "\n\n\n" not in result
        assert "Paragraph one." in result
        assert "Paragraph two." in result

    def test_handles_empty_string(self):
        assert _clean("") == ""

    def test_strips_trailing_whitespace_per_line(self):
        text = "line one   \nline two  "
        result = _clean(text)
        for line in result.splitlines():
            assert line == line.rstrip()

    def test_unicode_normalization(self):
        # Smart quotes and apostrophes should survive as readable unicode
        text = "\u201cThis is quoted\u201d and it\u2019s fine"
        result = _clean(text)
        assert len(result) > 0
        assert "fine" in result


# ---------------------------------------------------------------------------
# TXT / MD extraction tests
# ---------------------------------------------------------------------------

class TestTextFileExtraction:

    def test_basic_txt(self, tmp_path):
        p = tmp_path / "policy.txt"
        p.write_text("Access Control Policy\n\nRole-based access control is enforced.", encoding="utf-8")
        result = extract_text(p)
        assert "role-based access" in result.lower()

    def test_page_number_cleaned_from_txt(self, tmp_path):
        p = tmp_path / "policy.txt"
        p.write_text("Access policy content.\n\nPage 1\n\nMore content.", encoding="utf-8")
        result = extract_text(p)
        assert "Page 1" not in result
        assert "Access policy content." in result

    def test_markdown_file(self, tmp_path):
        p = tmp_path / "policy.md"
        p.write_text("# Access Control Policy\n\nRBAC is enforced for all users.", encoding="utf-8")
        result = extract_text(p)
        assert "rbac" in result.lower()

    def test_accepts_string_path(self, tmp_path):
        p = tmp_path / "policy.txt"
        p.write_text("some content here to read", encoding="utf-8")
        # Pass as string, not Path object
        result = extract_text(str(p))
        assert "some content" in result

    def test_short_file_triggers_warning(self, tmp_path):
        p = tmp_path / "tiny.txt"
        p.write_text("Hi", encoding="utf-8")
        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            extract_text(p)
        assert any(issubclass(w.category, LowQualityExtractionWarning) for w in caught)


# ---------------------------------------------------------------------------
# DOCX extraction tests
# ---------------------------------------------------------------------------

class TestDocxExtraction:

    def test_extracts_paragraphs(self, tmp_path):
        doc = Document()
        doc.add_paragraph("MFA is enforced for all users via Okta.")
        doc.add_paragraph("Least privilege principle applies to all roles.")
        p = tmp_path / "evidence.docx"
        doc.save(str(p))

        result = extract_text(p)
        assert "mfa is enforced" in result.lower()
        assert "least privilege" in result.lower()

    def test_extracts_table_content(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Role"
        table.cell(0, 1).text = "Access Level"
        table.cell(1, 0).text = "Admin"
        table.cell(1, 1).text = "Full"
        p = tmp_path / "access_matrix.docx"
        doc.save(str(p))

        result = extract_text(p)
        assert "admin" in result.lower()
        assert "access level" in result.lower()

    def test_extracts_both_paragraphs_and_tables(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Access control policy overview.")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Engineer"
        table.cell(0, 1).text = "Read-only production"
        p = tmp_path / "combined.docx"
        doc.save(str(p))

        result = extract_text(p)
        assert "access control policy" in result.lower()
        assert "engineer" in result.lower()


# ---------------------------------------------------------------------------
# Error handling tests
# ---------------------------------------------------------------------------

class TestErrorHandling:

    def test_unsupported_format_raises(self, tmp_path):
        p = tmp_path / "spreadsheet.xlsx"
        p.write_bytes(b"fake xlsx content")
        with pytest.raises(ExtractionError, match="Unsupported file format"):
            extract_text(p)

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(ExtractionError, match="File not found"):
            extract_text(tmp_path / "does_not_exist.pdf")

    def test_error_message_includes_extension(self, tmp_path):
        p = tmp_path / "file.csv"
        p.write_text("col1,col2", encoding="utf-8")
        with pytest.raises(ExtractionError) as exc_info:
            extract_text(p)
        assert ".csv" in str(exc_info.value)